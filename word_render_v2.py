# word_render_v2.py

import argparse
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from openpyxl import load_workbook
from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# =========================================================
# CONFIG
# =========================================================

PLACEHOLDER_BODY = "[[REPORT_BODY]]"
PLACEHOLDER_APPENDIX = "[[APPENDIX_START]]"

STYLE_TITLE = "Title"
STYLE_SUBTITLE = "Subtitle"
STYLE_SECTION = "Section Heading"
STYLE_SUBSECTION = "Subheading"
STYLE_BODY = "Body Text"
STYLE_BULLET = "Bullet"
STYLE_CAPTION = "Caption"
STYLE_SOURCE = "Source Text"


# =========================================================
# HELPERS
# =========================================================

def clean_text(text: Any) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def normalize_paragraphs(text: str) -> List[str]:
    text = clean_text(text)
    if not text:
        return []
    parts = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    return parts if parts else [text]


def normalize_bullets(text_or_list: Union[str, List[str], None]) -> List[str]:
    if text_or_list is None:
        return []

    if isinstance(text_or_list, list):
        return [clean_text(x) for x in text_or_list if clean_text(x)]

    text = str(text_or_list).strip()
    if not text:
        return []

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    bullets = []

    for ln in lines:
        ln = re.sub(r"^[\-\u2022\*•]\s*", "", ln).strip()
        if ln:
            bullets.append(ln)

    return bullets


def read_json_report(report_json_path: str) -> Dict[str, Any]:
    with open(report_json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def read_text_report(report_txt_path: str) -> Dict[str, Any]:
    """
    Fallback parser for a plain-text report.

    Expected broad structure:
    SECTION TITLE
    content...

    This is a fallback only. JSON is preferred.
    """
    with open(report_txt_path, "r", encoding="utf-8") as f:
        raw = f.read()

    return {
        "raw_report": raw
    }


def load_report_payload(report_path: str) -> Dict[str, Any]:
    ext = Path(report_path).suffix.lower()

    if ext == ".json":
        return read_json_report(report_path)

    if ext in {".txt", ".md"}:
        return read_text_report(report_path)

    raise ValueError(f"Unsupported report file type: {ext}")


def read_links_from_excel(xlsx_path: str, sheet_name: Optional[str] = None) -> List[str]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb[wb.sheetnames[0]]

    links = []
    for row in ws.iter_rows(min_col=1, max_col=1):
        cell = row[0].value
        if cell is None:
            continue
        val = clean_text(cell)
        if val:
            links.append(val)

    return links


def paragraph_contains_placeholder(paragraph: Paragraph, placeholder: str) -> bool:
    return placeholder in paragraph.text


def replace_text_in_paragraph(paragraph: Paragraph, placeholder: str, replacement: str) -> None:
    if placeholder not in paragraph.text:
        return

    inline = paragraph.runs
    for run in inline:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)


def replace_placeholders_in_doc(doc: Document, context: Dict[str, Any]) -> None:
    """
    Replaces simple {{placeholder}} tokens in all paragraphs and table cells.
    """
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            token = f"{{{{{key}}}}}"
            if token in paragraph.text:
                replace_text_in_paragraph(paragraph, token, clean_text(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in context.items():
                        token = f"{{{{{key}}}}}"
                        if token in paragraph.text:
                            replace_text_in_paragraph(paragraph, token, clean_text(value))


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    """
    Insert a new paragraph directly after the given paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Any:
    """
    Insert an empty table after a given paragraph.
    """
    table = paragraph._parent.add_table(rows=rows, cols=cols)
    paragraph._element.addnext(table._element)
    return table


def find_placeholder_paragraph(doc: Document, placeholder: str) -> Optional[Paragraph]:
    for p in doc.paragraphs:
        if placeholder in p.text:
            return p
    return None


def set_cell_text(cell, text: str, style: Optional[str] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if style:
        p.style = style
    p.add_run(text)


def add_styled_paragraph(doc: Document, text: str, style: str = STYLE_BODY) -> Paragraph:
    p = doc.add_paragraph(style=style)
    p.add_run(clean_text(text))
    return p


def add_bullet_paragraph(doc: Document, text: str) -> Paragraph:
    p = doc.add_paragraph(style=STYLE_BULLET)
    p.add_run(clean_text(text))
    return p


def add_section_heading(doc: Document, text: str) -> Paragraph:
    return add_styled_paragraph(doc, text, STYLE_SECTION)


def add_subheading(doc: Document, text: str) -> Paragraph:
    return add_styled_paragraph(doc, text, STYLE_SUBSECTION)


def render_text_block(doc: Document, content: Union[str, List[str], Dict[str, Any], None]) -> None:
    """
    Render either:
    - plain string
    - list of strings
    - dict with paragraphs/bullets/title
    """
    if content is None:
        return

    if isinstance(content, dict):
        title = clean_text(content.get("title"))
        if title:
            add_subheading(doc, title)

        paragraphs = content.get("paragraphs", [])
        bullets = content.get("bullets", [])
        tables = content.get("tables", [])

        for para in paragraphs:
            para = clean_text(para)
            if para:
                add_styled_paragraph(doc, para, STYLE_BODY)

        for bullet in bullets:
            bullet = clean_text(bullet)
            if bullet:
                add_bullet_paragraph(doc, bullet)

        for tbl in tables:
            render_table(doc, tbl)

        return

    if isinstance(content, list):
        for item in content:
            item = clean_text(item)
            if item:
                add_styled_paragraph(doc, item, STYLE_BODY)
        return

    text = str(content).strip()
    if not text:
        return

    paragraphs = normalize_paragraphs(text)
    for para in paragraphs:
        if re.match(r"^[\-\u2022\*•]\s+", para):
            add_bullet_paragraph(doc, re.sub(r"^[\-\u2022\*•]\s+", "", para).strip())
        else:
            add_styled_paragraph(doc, para, STYLE_BODY)


def render_table(doc: Document, table_obj: Dict[str, Any]) -> None:
    """
    Expected table object:
    {
        "caption": "...",
        "headers": ["A","B","C"],
        "rows": [["1","2","3"], ...]
    }
    """
    caption = clean_text(table_obj.get("caption"))
    headers = table_obj.get("headers", [])
    rows = table_obj.get("rows", [])

    if caption:
        add_styled_paragraph(doc, caption, STYLE_CAPTION)

    if not headers and not rows:
        return

    ncols = max(len(headers), max((len(r) for r in rows), default=0))
    if ncols == 0:
        return

    table = doc.add_table(rows=1, cols=ncols)
    table.style = table_obj.get("style", "Table Grid")

    hdr_cells = table.rows[0].cells
    for i in range(ncols):
        text = clean_text(headers[i]) if i < len(headers) else ""
        set_cell_text(hdr_cells[i], text, STYLE_BODY)

    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            text = clean_text(row[i]) if i < len(row) else ""
            set_cell_text(cells[i], text, STYLE_BODY)


# =========================================================
# MAIN REPORT RENDERER
# =========================================================

def render_word_report(
    report_payload: Dict[str, Any],
    template_path: str,
    output_path: str,
    links_xlsx_path: Optional[str] = None,
    links_sheet_name: Optional[str] = None
) -> None:
    doc = Document(template_path)

    context = {
        "company_name": report_payload.get("company_name", ""),
        "short_name": report_payload.get("short_name", ""),
        "ticker": report_payload.get("ticker", ""),
        "report_date": report_payload.get("report_date", ""),
        "analyst_name": report_payload.get("analyst_name", ""),
        "disclaimer": report_payload.get("disclaimer", ""),
    }
    replace_placeholders_in_doc(doc, context)

    body_anchor = find_placeholder_paragraph(doc, PLACEHOLDER_BODY)
    if body_anchor is None:
        raise ValueError(
            f"Template must contain the placeholder paragraph {PLACEHOLDER_BODY}"
        )

    # Clear the placeholder text so it does not appear in the final report
    body_anchor.text = ""

    cursor = body_anchor

    # =====================================================
    # TITLE BLOCK IF TEMPLATE USES THE BODY SECTION FOR IT
    # =====================================================
    # If your template already has title-page fields, these are
    # typically handled by replace_placeholders_in_doc above.
    # The body begins after the anchor.

    # =====================================================
    # INVESTMENT THESIS
    # =====================================================
    thesis = report_payload.get("investment_thesis", {})
    cursor = insert_paragraph_after(cursor, "", None)
    cursor = insert_paragraph_after(cursor, "INVESTMENT THESIS", STYLE_SECTION)

    render_text_block_to_doc_after(cursor, thesis)

    # =====================================================
    # 5 KEY HIGHLIGHTS
    # =====================================================
    key_highlights = report_payload.get("key_highlights", [])
    cursor = insert_paragraph_after(find_last_paragraph(doc), "", None)
    cursor = insert_paragraph_after(cursor, "5 KEY HIGHLIGHTS", STYLE_SECTION)

    for idx, item in enumerate(key_highlights, 1):
        if isinstance(item, dict):
            title = clean_text(item.get("title", f"Highlight {idx}"))
            body = item.get("body", "")
            cursor = insert_paragraph_after(cursor, f"{idx}. {title}", STYLE_SUBSECTION)
            render_text_block_to_doc_after(cursor, body)
        else:
            cursor = insert_paragraph_after(cursor, f"{idx}. {clean_text(item)}", STYLE_BULLET)

    # =====================================================
    # COMPANY OVERVIEW
    # =====================================================
    company_overview = report_payload.get("company_overview", {})
    cursor = insert_paragraph_after(find_last_paragraph(doc), "", None)
    cursor = insert_paragraph_after(cursor, "COMPANY OVERVIEW", STYLE_SECTION)
    render_text_block_to_doc_after(cursor, company_overview)

    # =====================================================
    # HIGHLIGHT DEEP DIVES
    # =====================================================
    highlight_sections = report_payload.get("highlight_sections", [])
    if highlight_sections:
        cursor = insert_paragraph_after(find_last_paragraph(doc), "", None)
        cursor = insert_paragraph_after(cursor, "DETAILS ON THE 5 HIGHLIGHTS", STYLE_SECTION)

        for idx, item in enumerate(highlight_sections, 1):
            if isinstance(item, dict):
                title = clean_text(item.get("title", f"Highlight {idx}"))
                body = item.get("body", {})
                cursor = insert_paragraph_after(cursor, f"{idx}. {title}", STYLE_SUBSECTION)
                render_text_block_to_doc_after(cursor, body)
            else:
                cursor = insert_paragraph_after(cursor, f"{idx}. {clean_text(item)}", STYLE_BULLET)

    # =====================================================
    # FINANCIALS
    # =====================================================
    financials = report_payload.get("financials", {})
    cursor = insert_paragraph_after(find_last_paragraph(doc), "", None)
    cursor = insert_paragraph_after(cursor, "FINANCIALS", STYLE_SECTION)
    render_text_block_to_doc_after(cursor, financials)

    # =====================================================
    # RISKS
    # =====================================================
    risks = report_payload.get("risks", {})
    cursor = insert_paragraph_after(find_last_paragraph(doc), "", None)
    cursor = insert_paragraph_after(cursor, "RISKS", STYLE_SECTION)
    render_text_block_to_doc_after(cursor, risks)

    # =====================================================
    # SOURCES APPENDIX FROM EXCEL LINK LIST
    # =====================================================
    if links_xlsx_path:
        links = read_links_from_excel(links_xlsx_path, links_sheet_name)

        if links:
            anchor = find_placeholder_paragraph(doc, PLACEHOLDER_APPENDIX)
            if anchor is None:
                # If no placeholder exists, append at the end.
                anchor = find_last_paragraph(doc)
                anchor = insert_paragraph_after(anchor, "", None)
            else:
                anchor.text = ""

            cursor = insert_paragraph_after(anchor, "SOURCE LIST", STYLE_SECTION)
            cursor = insert_paragraph_after(cursor, "Secondary research links used during report preparation.", STYLE_BODY)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Source #"
            table.rows[0].cells[1].text = "URL"

            for idx, link in enumerate(links, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = link

    doc.save(output_path)
    print(f"[INFO] Word report saved -> {output_path}")


def find_last_paragraph(doc: Document) -> Paragraph:
    if doc.paragraphs:
        return doc.paragraphs[-1]
    return doc.add_paragraph()


def render_text_block_to_doc_after(anchor: Paragraph, content: Union[str, List[str], Dict[str, Any], None]) -> Paragraph:
    """
    Render content immediately after the given anchor paragraph.
    Returns the last inserted paragraph so subsequent content can continue from there.
    """
    if content is None:
        return anchor

    current = anchor

    if isinstance(content, dict):
        title = clean_text(content.get("title"))
        if title:
            current = insert_paragraph_after(current, title, STYLE_SUBSECTION)

        paragraphs = content.get("paragraphs", [])
        bullets = content.get("bullets", [])
        tables = content.get("tables", [])

        for para in paragraphs:
            para = clean_text(para)
            if para:
                current = insert_paragraph_after(current, para, STYLE_BODY)

        for bullet in bullets:
            bullet = clean_text(bullet)
            if bullet:
                current = insert_paragraph_after(current, bullet, STYLE_BULLET)

        for tbl in tables:
            current = insert_table_and_advance(current, tbl)

        return current

    if isinstance(content, list):
        for item in content:
            item = clean_text(item)
            if item:
                current = insert_paragraph_after(current, item, STYLE_BODY)
        return current

    text = str(content).strip()
    if not text:
        return current

    paragraphs = normalize_paragraphs(text)
    for para in paragraphs:
        if re.match(r"^[\-\u2022\*•]\s+", para):
            current = insert_paragraph_after(
                current,
                re.sub(r"^[\-\u2022\*•]\s+", "", para).strip(),
                STYLE_BULLET
            )
        else:
            current = insert_paragraph_after(current, para, STYLE_BODY)

    return current


def insert_table_and_advance(anchor: Paragraph, table_obj: Dict[str, Any]) -> Paragraph:
    """
    Insert a table after the anchor paragraph and return the paragraph immediately following it.
    """
    caption = clean_text(table_obj.get("caption"))
    headers = table_obj.get("headers", [])
    rows = table_obj.get("rows", [])

    current = anchor

    if caption:
        current = insert_paragraph_after(current, caption, STYLE_CAPTION)

    ncols = max(len(headers), max((len(r) for r in rows), default=0))
    if ncols <= 0:
        return current

    # Add the table at document end and then move it after current via XML
    doc = current._parent if isinstance(current._parent, Document) else current._parent.part.document
    table = doc.add_table(rows=1, cols=ncols)
    table.style = table_obj.get("style", "Table Grid")

    current._element.addnext(table._element)

    hdr_cells = table.rows[0].cells
    for i in range(ncols):
        hdr_cells[i].text = clean_text(headers[i]) if i < len(headers) else ""

    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            cells[i].text = clean_text(row[i]) if i < len(row) else ""

    # Return the paragraph after the table, so rendering can continue
    tail = OxmlElement("w:p")
    table._element.addnext(tail)
    return Paragraph(tail, current._parent)


# =========================================================
# CLI
# =========================================================

def main():
    parser = argparse.ArgumentParser(description="Render institutional report to Word from structured output.")

    parser.add_argument("--report", required=True, help="Path to report JSON or TXT")
    parser.add_argument("--template", required=True, help="Path to Word template .docx")
    parser.add_argument("--output", required=True, help="Path to final Word .docx")
    parser.add_argument("--links_xlsx", required=False, default=None, help="Optional Excel file containing source links")
    parser.add_argument("--links_sheet", required=False, default=None, help="Optional sheet name for source links")

    args = parser.parse_args()

    report_payload = load_report_payload(args.report)

    # If the report is plain text, keep rendering limited to a simple body.
    # Best practice is to emit JSON from the generation stage.
    if "raw_report" in report_payload:
        report_payload = {
            "company_name": report_payload.get("company_name", ""),
            "short_name": report_payload.get("short_name", ""),
            "ticker": report_payload.get("ticker", ""),
            "report_date": report_payload.get("report_date", ""),
            "analyst_name": report_payload.get("analyst_name", ""),
            "investment_thesis": {
                "paragraphs": normalize_paragraphs(report_payload["raw_report"])
            },
            "key_highlights": [],
            "company_overview": {},
            "highlight_sections": [],
            "financials": {},
            "risks": {}
        }

    render_word_report(
        report_payload=report_payload,
        template_path=args.template,
        output_path=args.output,
        links_xlsx_path=args.links_xlsx,
        links_sheet_name=args.links_sheet
    )


if __name__ == "__main__":
    main()